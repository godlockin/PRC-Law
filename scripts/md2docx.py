#!/usr/bin/env python3
"""
md2docx.py — Markdown → Word DOCX 转换器（律鉴格式）

将律鉴的 Markdown 报告转换为符合中国律师实务格式的 Word 文档。
支持律鉴 v8.6 demo 报告 + 任意中文 Markdown。

格式特性：
- 中文字体（宋体/黑体）
- 标题分级（一/二/三级）
- 表格（带边框）
- 代码块（等宽字体）
- 引用块（律师审阅闸）
- 列表（多级）

使用：
    from md2docx import convert
    convert("input.md", "output.docx", title="律鉴审查报告")

或 CLI：
    python3 md2docx.py input.md output.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# 中文字体配置
FONT_BODY = "SimSun"        # 宋体（正文）
FONT_HEADING = "SimHei"     # 黑体（标题）
FONT_MONO = "Consolas"        # 等宽字体（代码）
FONT_FANGSONG = "FangSong"  # 仿宋（合同/法律文书常用）

# 颜色
COLOR_HEADING = RGBColor(0x00, 0x00, 0x00)        # 黑色
COLOR_EMPHASIS = RGBColor(0xCC, 0x00, 0x00)      # 红色（强调）
COLOR_TABLE_HEADER = RGBColor(0xF2, 0xF2, 0xF2)  # 浅灰（表头背景）


def set_chinese_font(run, font_name: str = FONT_BODY, size: int = 11):
    """设置中文字体（同时设置西文字体和东亚字体）"""
    run.font.name = font_name
    # 关键：设置东亚字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    run.font.size = Pt(size)


def add_horizontal_line(paragraph):
    """添加水平线"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pPr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '12',
        qn('w:color'): '666666'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_table_row(line: str) -> list:
    """解析表格行"""
    line = line.strip()
    if line.startswith("|") and line.endswith("|"):
        line = line[1:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    """判断是否为表格分隔符行 |---|---|"""
    line = line.strip()
    return bool(re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", line))


def add_table(doc, rows, headers=None):
    """添加表格"""
    if not rows and not headers:
        return

    all_rows = (headers or []) + (rows or [])
    if not all_rows:
        return

    n_cols = max(len(r) for r in all_rows)
    table = doc.add_table(rows=len(all_rows), cols=n_cols)
    # 不使用预定义样式（避免中文竖切问题）
    table.style = "Table Grid"
    # 设置表格自动调整
    table.autofit = True

    for i, row_data in enumerate(all_rows):
        row = table.rows[i]
        for j in range(n_cols):
            cell = row.cells[j]
            text = row_data[j] if j < len(row_data) else ""
            # 清空 cell 现有内容（保留 paragraph 结构）
            for p in cell.paragraphs:
                # 移除所有现有 run
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            p = cell.paragraphs[0]
            # 清理 markdown 格式
            text_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text_clean = re.sub(r"\*(.+?)\*", r"\1", text_clean)
            text_clean = re.sub(r"`(.+?)`", r"\1", text_clean)
            run = p.add_run(text_clean)
            is_header = i == 0 and headers is not None
            set_chinese_font(
                run,
                font_name=FONT_HEADING if is_header else FONT_BODY,
                size=10,
            )
            if is_header:
                run.font.bold = True
                # 表头加底色
                shading_elm = cell._element.tcPr.makeelement(qn('w:shd'), {})
                shading_elm.set(qn('w:fill'), 'F2F2F2')
                cell._element.tcPr.append(shading_elm)


def parse_and_convert(md_text: str, title: str = "律鉴法律文书"):
    """解析 Markdown 并转换为 Word 文档"""
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    set_chinese_font(run, font_name=FONT_HEADING, size=18)
    run.font.bold = True

    # 副标题
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("（律鉴 PRC-Law 智能分析系统）")
    set_chinese_font(run, font_name=FONT_BODY, size=10)
    run.font.italic = True

    # 水平线
    add_horizontal_line(doc.add_paragraph())

    # 解析 Markdown
    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []
    table_header = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, font_name=FONT_MONO, size=9)
            i += 1
            continue

        # 表格
        if stripped.startswith("|") and "|" in stripped[1:]:
            if not in_table:
                # 检查下一行是否是分隔符
                if i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                    in_table = True
                    table_header = parse_table_row(line)
                    i += 2  # 跳过分隔符
                    continue
                else:
                    # 不是表格，跳过
                    i += 1
                    continue
            else:
                if stripped:
                    table_rows.append(parse_table_row(line))
                i += 1
                # 检查是否结束
                if i >= len(lines) or not (lines[i].strip().startswith("|") and "|" in lines[i].strip()[1:]):
                    add_table(doc, table_rows, table_header)
                    in_table = False
                    table_rows = []
                    table_header = None
                continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # 清理 markdown 格式
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)

            heading = doc.add_heading(text, level=min(level, 6))
            for run in heading.runs:
                sizes = {1: 16, 2: 14, 3: 12, 4: 11, 5: 10, 6: 10}
                set_chinese_font(run, font_name=FONT_HEADING, size=sizes.get(level, 11))
                run.font.bold = True
                run.font.color.rgb = COLOR_HEADING
            i += 1
            continue

        # 引用块（律师审阅闸）
        if stripped.startswith(">"):
            text = stripped[1:].strip()
            # 清理 markdown 格式
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.5)
            run = p.add_run("│ " + text)
            set_chinese_font(run, font_name=FONT_FANGSONG, size=10)
            run.font.italic = True
            # 左侧边框
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            left = pPr.makeelement(qn('w:left'), {
                qn('w:val'): 'single',
                qn('w:sz'): '24',
                qn('w:color'): 'CC0000',
                qn('w:space'): '4'
            })
            pBdr.append(left)
            pPr.append(pBdr)
            i += 1
            continue

        # 列表
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)$", line)
        if m:
            indent = len(m.group(1))
            marker = m.group(2)
            text = m.group(3)
            # 清理 markdown 格式
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)

            style_name = "List Number" if marker.endswith(".") else "List Bullet"
            p = doc.add_paragraph(text, style=style_name)
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.3)
            for run in p.runs:
                set_chinese_font(run, font_name=FONT_BODY, size=11)
            i += 1
            continue

        # 粗体
        is_horizontal_rule = re.match(r"^---+$", stripped)

        # 段落
        if stripped and not is_horizontal_rule:
            # 清理 markdown 格式
            p = doc.add_paragraph()
            # 简单的内联格式处理
            segments = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", stripped)
            for seg in segments:
                if not seg:
                    continue
                if seg.startswith("**") and seg.endswith("**"):
                    run = p.add_run(seg[2:-2])
                    run.font.bold = True
                elif seg.startswith("*") and seg.endswith("*"):
                    run = p.add_run(seg[1:-1])
                    run.font.italic = True
                elif seg.startswith("`") and seg.endswith("`"):
                    run = p.add_run(seg[1:-1])
                    set_chinese_font(run, font_name=FONT_MONO, size=10)
                else:
                    run = p.add_run(seg)
                set_chinese_font(run, font_name=FONT_BODY, size=11)
            i += 1
        elif is_horizontal_rule:
            add_horizontal_line(doc.add_paragraph())
            i += 1
        else:
            i += 1

    return doc


def convert(input_path: str, output_path: str, title: Optional[str] = None):
    """主转换函数"""
    md_text = Path(input_path).read_text(encoding="utf-8")
    if title is None:
        # 从 markdown 第一个 # 提取
        for line in md_text.split("\n"):
            m = re.match(r"^#\s+(.+)$", line.strip())
            if m:
                title = m.group(1)
                break
    if not title:
        title = "律鉴法律文书"

    doc = parse_and_convert(md_text, title)
    doc.save(output_path)
    print(f"✅ 已转换: {input_path} → {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 md2docx.py <input.md> <output.docx> [title]")
        sys.exit(1)
    title = sys.argv[3] if len(sys.argv) > 3 else None
    convert(sys.argv[1], sys.argv[2], title)